"""
============================================================
VaultMind — Doküman Yükleyici (Document Loader)
============================================================

BU DOSYA NE YAPIYOR?
--------------------
Farklı formatlardaki dosyaları (PDF, DOCX, Markdown, TXT) okuyup
tek bir ortak formata dönüştürür: düz metin + metadata.

NEDEN GEREKLİ?
--------------
Bir şirkette dokümanlar onlarca farklı formatta olabilir:
- İK politikaları: PDF
- Teknik kılavuzlar: Markdown
- Toplantı notları: DOCX
- Log dosyaları: TXT

RAG pipeline'ı sadece "metin" ile çalışır. Bu katman, tüm
formatları metne dönüştürerek pipeline'ın geri kalanının
dosya formatından bağımsız çalışmasını sağlar.

ÖĞRENME NOKTALARI:
-----------------
1. Strategy Pattern — Her dosya formatı için ayrı yükleyici
2. Dataclass — Yüklenen dokümanın yapısal temsili
3. Path handling — Python'da dosya yolu yönetimi
4. Error handling — Desteklenmeyen format ve hata yönetimi
"""

from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class LoadedDocument:
    """
    Yüklenen bir dokümanı temsil eder.

    Hangi formattan gelirse gelsin, sonuçta şu bilgileri taşır:
    - content: Dokümanın düz metin içeriği
    - metadata: Dosya hakkında ek bilgiler

    Metadata örneği:
    {
        "source": "ik_politikasi.pdf",
        "file_type": "pdf",
        "page_count": 5,
        "security_level": "internal",  # RBAC için (Phase 5)
    }
    """
    content: str                               # Dokümanın düz metni
    metadata: dict = field(default_factory=dict)  # Ek bilgiler


def load_txt(file_path: Path) -> LoadedDocument:
    """
    Düz metin (.txt) dosyasını yükler.

    En basit yükleyici — dosyayı olduğu gibi okur.
    """
    content = file_path.read_text(encoding="utf-8")
    return LoadedDocument(
        content=content,
        metadata={
            "source": file_path.name,
            "file_type": "txt",
            "file_path": str(file_path),
            "char_count": len(content),
        }
    )


def load_markdown(file_path: Path) -> LoadedDocument:
    """
    Markdown (.md) dosyasını yükler.

    Markdown zaten düz metin tabanlıdır, ekstra işlem gerekmez.
    Başlıklar (#), listeler (-), kalın (**) gibi işaretler
    metinle birlikte kalır — LLM bunları anlayabilir.
    """
    content = file_path.read_text(encoding="utf-8")
    return LoadedDocument(
        content=content,
        metadata={
            "source": file_path.name,
            "file_type": "markdown",
            "file_path": str(file_path),
            "char_count": len(content),
        }
    )


def load_pdf(file_path: Path) -> LoadedDocument:
    """
    PDF dosyasını yükler.

    pypdf kütüphanesi ile PDF'in her sayfasını okur ve
    tek bir metin haline getirir.

    Öğrenme notu:
    PDF'ler görsel formatta veri tutar (fontlar, konumlar, görseller).
    Metin çıkarma her zaman mükemmel olmayabilir — özellikle taranmış
    (scanned) PDF'lerde OCR gerekebilir. Şimdilik metin tabanlı
    PDF'lerle çalışıyoruz.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(text)

    content = "\n\n".join(pages)

    return LoadedDocument(
        content=content,
        metadata={
            "source": file_path.name,
            "file_type": "pdf",
            "file_path": str(file_path),
            "page_count": len(reader.pages),
            "char_count": len(content),
        }
    )


def load_docx(file_path: Path) -> LoadedDocument:
    """
    Word (.docx) dosyasını yükler.

    python-docx kütüphanesi ile DOCX'in paragraflarını okur.
    Tablolar ve görseller şimdilik atlanır — sadece metin paragrafları.
    """
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    content = "\n\n".join(paragraphs)

    return LoadedDocument(
        content=content,
        metadata={
            "source": file_path.name,
            "file_type": "docx",
            "file_path": str(file_path),
            "paragraph_count": len(paragraphs),
            "char_count": len(content),
        }
    )


# ============================================================
# Format → Yükleyici eşlemesi
# ============================================================
# Desteklenen dosya uzantıları ve karşılık gelen yükleyici fonksiyonları.
# Yeni bir format eklemek için buraya bir satır eklemek yeterlidir.

LOADERS = {
    ".txt": load_txt,
    ".md": load_markdown,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(file_path: str | Path) -> LoadedDocument:
    """
    Herhangi bir desteklenen formattaki dokümanı yükler.

    Dosya uzantısına bakarak doğru yükleyiciyi otomatik seçer.
    Bu, Factory Pattern'in basit bir uygulamasıdır.

    Args:
        file_path: Yüklenecek dosyanın yolu

    Returns:
        LoadedDocument: Düz metin + metadata

    Raises:
        FileNotFoundError: Dosya bulunamadığında
        ValueError: Desteklenmeyen dosya formatında

    Kullanım:
        doc = load_document("sample_docs/ik_politikasi.md")
        print(doc.content[:100])
        print(doc.metadata)
    """
    path = Path(file_path)

    # Dosya var mı kontrol et
    if not path.exists():
        raise FileNotFoundError(f"Dosya bulunamadi: {path}")

    # Uzantıya göre doğru yükleyiciyi seç
    suffix = path.suffix.lower()
    loader = LOADERS.get(suffix)

    if loader is None:
        supported = ", ".join(LOADERS.keys())
        raise ValueError(
            f"Desteklenmeyen dosya formati: '{suffix}'\n"
            f"Desteklenen formatlar: {supported}"
        )

    return loader(path)


def load_directory(dir_path: str | Path) -> list[LoadedDocument]:
    """
    Bir klasördeki tüm desteklenen dokümanları yükler.

    Alt klasörlere de bakar (recursive).

    Args:
        dir_path: Taranacak klasör yolu

    Returns:
        list[LoadedDocument]: Yüklenen dokümanların listesi

    Kullanım:
        docs = load_directory("sample_docs/")
        print(f"{len(docs)} dokuman yuklendi")
    """
    path = Path(dir_path)

    if not path.is_dir():
        raise NotADirectoryError(f"Klasor bulunamadi: {path}")

    documents = []
    supported_extensions = set(LOADERS.keys())

    # Tüm dosyaları recursive olarak tara
    for file_path in sorted(path.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                doc = load_document(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"  [WARN] {file_path.name} yuklenemedi: {e}")

    return documents
